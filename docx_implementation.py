from docx.shared import Inches, Cm, Pt
import sys
from docx import Document
import subprocess
import os

p_name = sys.argv[1]
email = sys.argv[2]
flight = sys.argv[3]
departure_date = sys.argv[4]
departure_time = sys.argv[5]
source = sys.argv[6]
destination = sys.argv[7]
total_seats = sys.argv[8]

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Arial Nova'
font.size = Pt(10)

document.add_picture('AIMAirwaysLogo.png', width=Inches(1.5))

document.add_heading('Booking Confirmation', 0)

records = (('Passenger Name', p_name),
           ('Email', email),
           ('Flight Number', flight),
           ('Departure Date', departure_date),
           ('Departure Time', departure_time),
           ('Start location', source),
           ('Destination location', destination),
           ('Total seats booked', total_seats))

output_filename = f'{p_name}_receipt.docx'

table = document.add_table(rows=1, cols=2)
table.style = 'LightShading-Accent1'  # OR TableGrid
hdr_cells = table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('Particulars').bold = True
hdr_cells[1].paragraphs[0].add_run('Values').bold = True

for particular, value in records:
    row_cells = table.add_row().cells
    row_cells[0].paragraphs[0].add_run(particular).bold = True
    row_cells[1].text = str(value)

for row in table.rows:
    row.height = Cm(0.7)

document.save(output_filename)

process = subprocess.Popen(['docx2pdf', output_filename, os.path.join(os.getcwd(), output_filename.split('.docx')[0] + '.pdf')],
                           stdout=subprocess.PIPE,
                           stderr=subprocess.PIPE)
stdout, stderr = process.communicate()
print(stdout)
print(stderr)
